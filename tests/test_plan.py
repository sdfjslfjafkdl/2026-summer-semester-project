"""투자계획서 작성 지원 테스트.

conftest 가 LLM 을 강제로 끄므로 이 테스트는 네트워크 없이 돈다.
즉 여기서 통과한다는 것은 "키 없이도 초안과 docx 가 나온다"는 뜻이기도 하다.
"""

from __future__ import annotations

import io

import pytest
from docx import Document
from docx.oxml.ns import qn

from app.data.plan_sections import (
    ASSISTED_SECTIONS,
    AUTO_SECTIONS,
    CHAPTERS,
    FORMAT_RULES,
    MANUAL_SECTIONS,
    SECTIONS,
)
from app.services import plan_store


@pytest.fixture(autouse=True)
def isolated_plans(tmp_path, monkeypatch):
    """계획서 저장소를 테스트마다 비운다."""
    from app.config import get_settings

    monkeypatch.setenv("PLAN_DIR", str(tmp_path / "plans"))
    get_settings.cache_clear()
    plan_store.reset_for_tests()
    yield
    plan_store.reset_for_tests()
    get_settings.cache_clear()


@pytest.fixture
def plan_id(client) -> str:
    response = client.post("/api/plan/draft", json={"region": "제천시", "year": 2026})
    assert response.status_code == 200, response.text
    return response.json()["data"]["plan_id"]


# ── 1. 레지스트리와 서식 목차 일치 ──────────────────────────────


def test_registry_covers_every_chapter_of_the_template():
    """서식 목차는 Ⅰ~Ⅵ 여섯 장이다. 없는 장을 만들지도, 빠뜨리지도 않는다."""
    assert set(CHAPTERS) == {"Ⅰ", "Ⅱ", "Ⅲ", "Ⅳ", "Ⅴ", "Ⅵ"}
    assert {s.chapter for s in SECTIONS} == set(CHAPTERS)
    assert CHAPTERS["Ⅰ"] == "지역 여건분석 및 전망"
    assert CHAPTERS["Ⅵ"] == "기타"


def test_every_section_carries_the_guides_from_the_template():
    for section in SECTIONS:
        assert section.writing_guide, f"{section.number} 에 【작성내용】이 없습니다"
        assert section.evaluation_focus, f"{section.number} 에 【기술 방향과 평가의 주안점】이 없습니다"
        assert 10 <= section.source_page <= 38, f"{section.number} 의 안내서 쪽 번호가 이상합니다"
        assert section.fill_mode in {"auto", "assisted", "manual"}


def test_fill_modes_match_the_agreed_split():
    assert set(AUTO_SECTIONS) == {"1-1", "6-2-1"}
    assert set(ASSISTED_SECTIONS) == {"1-2", "3-1", "3-3", "3-4-1", "3-annex"}
    # Ⅱ, Ⅳ, Ⅴ 는 통째로 사람 몫이다
    for section in SECTIONS:
        if section.chapter in {"Ⅱ", "Ⅳ", "Ⅴ"}:
            assert section.fill_mode == "manual", f"{section.number} 는 manual 이어야 합니다"
    for section_id in ("3-2", "3-4-5", "3-5", "6-1"):
        assert section_id in MANUAL_SECTIONS


def test_goal_section_follows_the_guides_warning_about_execution_rate():
    """안내서는 단순 실적지표(예산 집행률)를 목표로 쓰지 말라고 한다."""
    goal = next(s for s in SECTIONS if s.section_id == "3-3")
    assert "단순 실적지표" in goal.reference_note
    assert "예산 집행률" in goal.reference_note


def test_format_rules_come_from_the_template():
    assert FORMAT_RULES["body_font"] == "휴먼명조"
    assert FORMAT_RULES["body_size_pt"] == 15
    assert FORMAT_RULES["note_font"] == "중고딕"
    assert FORMAT_RULES["note_size_pt"] == 13
    assert FORMAT_RULES["margin_top_mm"] == 15 and FORMAT_RULES["margin_left_mm"] == 20
    assert FORMAT_RULES["header_mm"] == 10 and FORMAT_RULES["page_numbers"] is True


# ── 2. auto 섹션 채움 ─────────────────────────────────────────


def test_draft_fills_auto_sections_from_layer1(client, plan_id, panel):
    data = client.get(f"/api/plan/{plan_id}").json()["data"]
    sections = {s["section_id"]: s for s in data["sections"]}

    population = sections["1-1"]
    assert population["status"] == "filled"
    assert population["source"] == "layer1"
    assert "인구현황" in population["content"]
    assert "청년 순이동률 추이" in population["content"]

    yearly_mean = panel.df[(panel.df.region == "제천시") & (panel.df.year == 2024)][
        "population_total"
    ].mean()
    assert f"{yearly_mean:,.0f}" in population["content"]

    fund = sections["6-2-1"]
    assert fund["status"] == "filled"
    frame = panel.fund_year_frame(2024)
    row = frame[frame.region == "제천시"].iloc[0]
    assert f"{row['fund_allocation_million_krw']:,.0f}" in fund["content"]


def test_draft_records_called_endpoints_and_progress(client, plan_id):
    data = client.get(f"/api/plan/{plan_id}").json()["data"]
    endpoints = set(data["called_endpoints"])
    assert "/api/funds/local-extinction/regions" in endpoints
    assert "/api/funds/local-extinction/trend" in endpoints
    assert "/api/panel/timeseries" in endpoints
    assert "/api/proposal" in endpoints
    assert "/api/evidence/projects" in endpoints

    progress = data["progress"]
    assert progress["total_sections"] == len(SECTIONS)
    assert progress["auto_filled"] == 2
    assert progress["awaiting_human"] == progress["assisted_pending"] + progress["manual_pending"]
    assert progress["filled"] + progress["awaiting_human"] == progress["total_sections"]


def test_manual_sections_stay_empty_but_carry_guidance(client, plan_id):
    sections = {s["section_id"]: s for s in client.get(f"/api/plan/{plan_id}").json()["data"]["sections"]}
    for section_id in MANUAL_SECTIONS:
        section = sections[section_id]
        assert section["content"] is None, f"{section['number']} 를 서버가 채웠습니다"
        assert section["status"] == "placeholder"
        assert section["guidance"]["writing_guide"]


def test_goal_targets_are_derived_from_the_panel(client, plan_id, panel):
    section = next(
        s for s in client.get(f"/api/plan/{plan_id}").json()["data"]["sections"] if s["section_id"] == "3-3"
    )
    assert "youth_net_migration_rate_per_1000" in section["content"]
    assert "집행률은 투입 진행률이므로 성과지표로 쓰지 않는다" in section["content"]
    assert "연차별 목표값 후보" in section["content"]

    labels = {p["label"]: p for p in section["data_points"]}
    for label in ("1차년 목표", "2차년 목표", "3차년 목표"):
        assert label in labels
        assert labels[label]["source_endpoint"] == "internal:plan_goal_targets"

    # 최근 실적은 패널 값 그대로여야 한다
    recent = panel.df[(panel.df.region == "제천시") & (panel.df.year == 2024)][
        "youth_net_migration_rate_per_1000"
    ].mean()
    assert f"{recent:,.2f}" in section["content"]


def test_draft_never_claims_causal_effect(client, plan_id):
    data = client.get(f"/api/plan/{plan_id}").json()
    blob = " ".join(s["content"] or "" for s in data["data"]["sections"])
    for forbidden in ("효과가 입증", "효과가 있었습니다", "인과관계가 확인"):
        assert forbidden not in blob
    assert any("인과효과" in note for note in data["meta"]["notes"])


# ── 3. 사람 입력 반영 ────────────────────────────────────────


def test_manual_section_stores_human_input_verbatim(client, plan_id):
    text = "부지는 제천시 신월동 934-2번지 시유지이며 2025년 12월 확보를 마쳤다. 민원은 없다."
    response = client.post(f"/api/plan/{plan_id}/sections/3-2", json={"content": text})
    assert response.status_code == 200
    section = response.json()["data"]["section"]
    assert section["content"] == text
    assert section["source"] == "human_input"
    assert section["status"] == "filled"

    # 사람이 적은 값은 출처가 사람 입력으로 남는다
    with_values = client.post(
        f"/api/plan/{plan_id}/sections/3-2",
        json={"content": text, "values": {"총사업비": 20000}},
    ).json()["data"]["section"]
    sources = {p["source_endpoint"] for p in with_values["data_points"]}
    assert "human_input" in sources


def test_human_input_shows_up_in_the_plan_and_progress(client, plan_id):
    before = client.get(f"/api/plan/{plan_id}").json()["data"]["progress"]["filled"]
    client.post(f"/api/plan/{plan_id}/sections/2-1", json={"content": "주민 설명회 2회 개최"})
    after = client.get(f"/api/plan/{plan_id}").json()
    assert after["data"]["progress"]["filled"] == before + 1
    section = next(s for s in after["data"]["sections"] if s["section_id"] == "2-1")
    assert "주민 설명회" in section["content"]
    assert any(entry["action"] == "section_update" for entry in after["data"]["history"])


def test_assisted_section_composes_from_human_values(client, plan_id):
    response = client.post(
        f"/api/plan/{plan_id}/sections/3-annex",
        json={"content": "1번 부, 2번 부", "values": {"기금사업명": "청년 정착 지원사업"}},
    )
    section = response.json()["data"]["section"]
    assert "청년 정착 지원사업" in section["content"]
    assert section["source"] in {"template", "llm"}


# ── 4. 수치 가드 ────────────────────────────────────────────


def test_generated_sentence_with_unknown_number_is_dropped(client, plan_id, monkeypatch):
    """서버가 만든 문장에 근거 없는 숫자가 있으면 그 문장을 버린다."""
    from app.services import plan_llm

    monkeypatch.setattr(
        plan_llm,
        "compose_assisted",
        lambda *args, **kwargs: (
            "○ 제천시 청년 순이동률은 -77.7명/천명이다.\n○ 사업 추진에는 문제가 없다."
        ),
    )
    response = client.post(f"/api/plan/{plan_id}/sections/1-2", json={"content": "요약해줘"})
    data = response.json()["data"]

    assert "-77.7" not in (data["section"]["content"] or "")
    assert "77.7" in " ".join(data["rejected_numbers"])
    assert "사업 추진에는 문제가 없다" in data["section"]["content"]
    assert any("제거" in w for w in data["section"]["warnings"])


def test_numbers_typed_by_a_human_are_allowed_but_marked(client, plan_id, monkeypatch):
    """사람이 직접 넣은 값은 예외로 허용하되 출처를 사람 입력으로 표시한다."""
    from app.services import plan_llm

    monkeypatch.setattr(
        plan_llm, "compose_assisted", lambda *args, **kwargs: "○ 총사업비는 20000백만원이다."
    )
    response = client.post(
        f"/api/plan/{plan_id}/sections/3-1",
        json={"content": "총사업비 20000", "values": {"총사업비": 20000}},
    )
    section = response.json()["data"]["section"]
    assert "20000" in section["content"]
    human = [p for p in section["data_points"] if p["source_endpoint"] == "human_input"]
    assert any(p["value"] == 20000 for p in human)


def test_auto_section_numbers_are_traceable_to_endpoints(client, plan_id):
    sections = {s["section_id"]: s for s in client.get(f"/api/plan/{plan_id}").json()["data"]["sections"]}
    for section_id in ("1-1", "6-2-1"):
        points = sections[section_id]["data_points"]
        assert points
        assert all(p["source_endpoint"].startswith(("/api/", "internal:")) for p in points)


# ── 5. Ⅵ-2 소계 대조 ────────────────────────────────────────


def test_subtotal_mismatch_raises_a_warning(client, plan_id, panel):
    frame = panel.fund_year_frame(2022)
    actual = float(frame[frame.region == "제천시"]["fund_allocation_million_krw"].iloc[0])

    response = client.post(
        f"/api/plan/{plan_id}/sections/6-2-1",
        json={"content": "사업별 행", "values": {"2022_배분액": actual - 800}},
    )
    warnings = response.json()["data"]["section"]["warnings"]
    assert any("불일치" in w for w in warnings), warnings
    assert any(f"{actual:,.0f}" in w for w in warnings)


def test_matching_subtotal_raises_no_warning(client, plan_id, panel):
    frame = panel.fund_year_frame(2022)
    actual = float(frame[frame.region == "제천시"]["fund_allocation_million_krw"].iloc[0])

    response = client.post(
        f"/api/plan/{plan_id}/sections/6-2-1",
        json={"content": "사업별 행", "values": {"2022_배분액": actual}},
    )
    assert not any("불일치" in w for w in response.json()["data"]["section"]["warnings"])


def test_auto_subtotal_section_declares_the_manual_remainder(client, plan_id):
    section = next(
        s for s in client.get(f"/api/plan/{plan_id}").json()["data"]["sections"] if s["section_id"] == "6-2-1"
    )
    assert "사업 단위로 분해할 수 없" in section["manual_remainder"]
    assert "사업별 행" in section["content"]


# ── 6. 수정과 diff ──────────────────────────────────────────


def test_revise_without_a_target_asks_back(client, plan_id):
    response = client.post(f"/api/plan/{plan_id}/revise", json={"instruction": "좀 더 잘 써줘"})
    data = response.json()["data"]
    assert data["resolved"] is False
    assert data["clarification_question"]
    assert data["changed_sections"] == []
    assert data["candidate_sections"]


def test_revise_targets_only_the_named_section(client, plan_id):
    before = client.get(f"/api/plan/{plan_id}").json()["data"]
    response = client.post(
        f"/api/plan/{plan_id}/revise", json={"instruction": "Ⅲ-3 사업 목표를 보수적으로 조정"}
    )
    data = response.json()["data"]
    assert data["resolved"] is True
    assert data["changed_sections"] == ["3-3"]

    after = client.get(f"/api/plan/{plan_id}").json()["data"]
    untouched = {s["section_id"]: s["content"] for s in before["sections"] if s["section_id"] != "3-3"}
    for section in after["sections"]:
        if section["section_id"] != "3-3":
            assert section["content"] == untouched[section["section_id"]]


def test_revise_diff_matches_before_and_after(client, plan_id):
    data = client.post(
        f"/api/plan/{plan_id}/revise", json={"instruction": "Ⅲ-3 사업 목표 조정"}
    ).json()["data"]
    change = data["changes"][0]
    assert change["summary"]
    assert change["before"] != change["after"]
    assert change["diff"][0].startswith("--- 3-3")
    assert change["diff"][1].startswith("+++ 3-3")

    added = [line[1:] for line in change["diff"] if line.startswith("+") and not line.startswith("+++")]
    for line in added:
        assert line in (change["after"] or "")


# ── 7. 요약 ────────────────────────────────────────────────


def test_summary_is_short_and_flags_missing_sections(client, plan_id):
    data = client.get(f"/api/plan/{plan_id}/summary").json()["data"]
    assert 1 <= len(data["summary_sentences"]) <= 5
    assert data["is_submittable"] is False
    assert data["missing_required_sections"]
    assert any("초안" in s for s in data["summary_sentences"])
    assert any("참고안" in s or "인과효과가 아니" in s for s in data["summary_sentences"])
    assert all(p["source_endpoint"] for p in data["evidence"])


# ── 8. 내보내기 ────────────────────────────────────────────


def _docx(client, plan_id) -> Document:
    response = client.post(f"/api/plan/{plan_id}/export", json={"format": "docx"})
    assert response.status_code == 200
    assert "wordprocessingml" in response.headers["content-type"]
    return Document(io.BytesIO(response.content))


def test_docx_follows_the_template_table_of_contents(client, plan_id):
    document = _docx(client, plan_id)
    texts = [p.text for p in document.paragraphs]
    for numeral, title in CHAPTERS.items():
        assert f"{numeral}. {title}" in texts, f"{numeral}장이 문서에 없습니다"


def test_docx_marks_manual_sections_for_the_officer(client, plan_id):
    document = _docx(client, plan_id)
    texts = [p.text for p in document.paragraphs]
    assert sum(1 for t in texts if "[담당자 작성 필요]" in t) >= len(MANUAL_SECTIONS)

    # 안내서의 작성내용과 평가 주안점이 그 자리에 들어간다
    blob = "\n".join(texts)
    assert "【작성내용】" in blob
    assert "【기술 방향과 평가의 주안점】" in blob

    shaded = [
        p
        for p in document.paragraphs
        if p._p.find(qn("w:pPr")) is not None and p._p.find(qn("w:pPr")).find(qn("w:shd")) is not None
    ]
    assert shaded, "manual 구획에 회색 음영이 없습니다"


def test_docx_keeps_korean_intact_and_applies_format_rules(client, plan_id):
    document = _docx(client, plan_id)
    texts = [p.text for p in document.paragraphs if p.text.strip()]
    assert any(any("가" <= ch <= "힣" for ch in t) for t in texts)
    assert not any("�" in t for t in texts), "한글이 깨졌습니다"

    section = document.sections[0]
    assert round(section.top_margin.mm) == FORMAT_RULES["margin_top_mm"]
    assert round(section.left_margin.mm) == FORMAT_RULES["margin_left_mm"]
    assert round(section.header_distance.mm) == FORMAT_RULES["header_mm"]
    assert "PAGE" in section.footer.paragraphs[0]._p.xml, "쪽 번호 필드가 없습니다"


def test_docx_says_it_is_a_draft_and_needs_hwp(client, plan_id):
    blob = "\n".join(p.text for p in _docx(client, plan_id).paragraphs)
    assert "이 문서는 초안입니다" in blob
    assert "hwp" in blob


def test_export_reports_font_substitution(client, plan_id):
    response = client.post(f"/api/plan/{plan_id}/export", json={"format": "docx"})
    from urllib.parse import unquote

    notes = unquote(response.headers.get("X-Plan-Notes", ""))
    # 휴먼명조가 설치된 환경이면 메모가 없을 수 있다
    if notes:
        assert "휴먼명조" in notes or "중고딕" in notes


def test_pdf_export_keeps_korean(client, plan_id):
    from pypdf import PdfReader

    response = client.post(f"/api/plan/{plan_id}/export", json={"format": "pdf"})
    assert response.status_code == 200
    assert response.headers["content-type"] == "application/pdf"

    reader = PdfReader(io.BytesIO(response.content))
    text = "\n".join(page.extract_text() or "" for page in reader.pages)
    assert any("가" <= ch <= "힣" for ch in text)
    assert "지방소멸대응기금" in text


# ── 9. LLM 비활성 폴백 ──────────────────────────────────────


def test_everything_works_without_llm(client, plan_id):
    """키 없이도 초안·사람 입력·docx 가 나온다. conftest 가 LLM 을 끈 상태다."""
    from app.config import get_settings

    assert get_settings().llm_active is False

    sections = {s["section_id"]: s for s in client.get(f"/api/plan/{plan_id}").json()["data"]["sections"]}
    assert sections["1-1"]["status"] == "filled"  # auto 는 그대로 채워진다
    assert sections["6-2-1"]["status"] == "filled"
    assert sections["2-1"]["content"] is None  # manual 구획도 그대로 생성된다

    updated = client.post(
        f"/api/plan/{plan_id}/sections/1-2", json={"content": "청년 유출이 지속되고 있다"}
    ).json()["data"]["section"]
    assert updated["source"] == "template"  # LLM 대신 템플릿 문장
    assert "청년 유출이 지속되고 있다" in updated["content"]

    assert _docx(client, plan_id)


def test_revise_without_llm_keeps_the_instruction_visible(client, plan_id):
    data = client.post(
        f"/api/plan/{plan_id}/revise", json={"instruction": "Ⅲ-3 사업 목표를 더 보수적으로"}
    ).json()["data"]
    assert data["resolved"] is True
    change = data["changes"][0]
    assert "담당자 수정 지시" in (change["after"] or "")
    assert "LLM" in change["summary"]


# ── 10. 입력 검증 ──────────────────────────────────────────


def test_unknown_region_and_past_year_are_rejected(client):
    bad_region = client.post("/api/plan/draft", json={"region": "서울시", "year": 2026})
    assert bad_region.status_code == 404
    assert bad_region.json()["error"]["code"] == "unknown_region"

    past = client.post("/api/plan/draft", json={"region": "제천시", "year": 2024})
    assert past.status_code == 422
    assert past.json()["error"]["code"] == "invalid_plan_year"


def test_unknown_plan_and_section_return_allowed_values(client, plan_id):
    missing = client.get("/api/plan/plan_없음_2026_01")
    assert missing.status_code == 404
    assert missing.json()["error"]["code"] == "unknown_plan"

    bad_section = client.post(
        f"/api/plan/{plan_id}/sections/9-9", json={"content": "내용"}
    )
    assert bad_section.status_code == 404
    error = bad_section.json()["error"]
    assert error["code"] == "unknown_section"
    assert "1-1" in error["allowed_values"]
