"""투자계획서 초안 내보내기 (docx / pdf).

안내서 p.3 의 형식 규정을 따른다.
  본문 휴먼명조 15pt / 참고사항 중고딕 13pt
  용지여백 위·아래 15, 좌·우 20, 머리말·꼬리말 10 (mm)
  쪽 번호 기재

휴먼명조·중고딕은 한글(HWP) 계열 글꼴이라 대부분의 서버에 없다. 없으면 대체 글꼴을 쓰고
그 사실을 로그와 응답 notes 에 남긴다. 실제 제출은 hwp 서식으로 변환해야 하므로
문서 첫 장에 그 사실을 적는다.
"""

from __future__ import annotations

import io
import logging
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor

from app.data.plan_sections import CHAPTERS, FORMAT_RULES, SECTIONS, TEMPLATE_SOURCE
from app.services.plan_store import StoredPlan

logger = logging.getLogger(__name__)

MANUAL_SHADING = "D9D9D9"
MANUAL_MARKER = "[담당자 작성 필요]"

# 휴먼명조가 없을 때 쓸 명조 계열, 중고딕이 없을 때 쓸 고딕 계열
SERIF_FALLBACKS = ("바탕", "Batang", "AppleMyungjo", "Noto Serif KR", "Source Han Serif K")
SANS_FALLBACKS = ("맑은 고딕", "Malgun Gothic", "AppleGothic", "Noto Sans KR", "Apple SD Gothic Neo")

_FONT_DIRS = (
    "/System/Library/Fonts",
    "/Library/Fonts",
    str(Path.home() / "Library/Fonts"),
    "/usr/share/fonts",
    "/usr/local/share/fonts",
)


def _installed_fonts() -> set[str]:
    names: set[str] = set()
    for directory in _FONT_DIRS:
        path = Path(directory)
        if not path.exists():
            continue
        try:
            for file in path.rglob("*"):
                if file.suffix.lower() in {".ttf", ".otf", ".ttc", ".dfont"}:
                    names.add(file.stem)
        except OSError:
            continue
    return names


def resolve_fonts() -> tuple[str, str, list[str]]:
    """(본문 글꼴, 참고사항 글꼴, 안내 메모)."""
    installed = _installed_fonts()
    notes: list[str] = []

    body = str(FORMAT_RULES["body_font"])
    note_font = str(FORMAT_RULES["note_font"])

    if not any(body in name or name in body for name in installed):
        replacement = next((f for f in SERIF_FALLBACKS if f in installed), SERIF_FALLBACKS[0])
        notes.append(
            f"본문 글꼴 '{body}'가 이 환경에 없어 '{replacement}'로 대체했다. "
            "hwp 로 변환할 때 서식 규정대로 휴먼명조 15pt 로 맞춰야 한다."
        )
        logger.warning("본문 글꼴 %s 없음 → %s 로 대체", body, replacement)
        body = replacement

    if not any(note_font in name or name in note_font for name in installed):
        replacement = next((f for f in SANS_FALLBACKS if f in installed), SANS_FALLBACKS[0])
        notes.append(
            f"참고사항 글꼴 '{note_font}'가 이 환경에 없어 '{replacement}'로 대체했다."
        )
        logger.warning("참고사항 글꼴 %s 없음 → %s 로 대체", note_font, replacement)
        note_font = replacement

    return body, note_font, notes


def _set_font(run, name: str, size_pt: int, *, bold: bool = False, color: str | None = None) -> None:  # noqa: ANN001
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    # 한글은 eastAsia 글꼴을 따로 지정해야 의도한 서체로 나온다.
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rfonts.set(qn(attr), name)


def _shade(paragraph, fill: str) -> None:  # noqa: ANN001
    ppr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    ppr.append(shd)


def _add_page_number_footer(section, font: str) -> None:  # noqa: ANN001
    """쪽 번호는 필드 코드로 넣는다. 안내서가 모든 문서에 쪽 번호를 요구한다."""
    paragraph = section.footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    _set_font(run, font, 11)

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._element.append(begin)
    run._element.append(instr)
    run._element.append(end)


def _paragraph(document, text: str, font: str, size: int, *, bold: bool = False,
               shade: str | None = None, color: str | None = None, space_after: int = 4):  # noqa: ANN001
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.line_spacing = 1.6  # 안내서: 줄간격 160%
    if shade:
        _shade(paragraph, shade)
    run = paragraph.add_run(text)
    _set_font(run, font, size, bold=bold, color=color)
    return paragraph


def build_docx(plan: StoredPlan) -> tuple[bytes, list[str]]:
    body_font, note_font, notes = resolve_fonts()
    body_size = int(FORMAT_RULES["body_size_pt"])
    note_size = int(FORMAT_RULES["note_size_pt"])

    document = Document()
    section = document.sections[0]
    section.top_margin = Mm(int(FORMAT_RULES["margin_top_mm"]))
    section.bottom_margin = Mm(int(FORMAT_RULES["margin_bottom_mm"]))
    section.left_margin = Mm(int(FORMAT_RULES["margin_left_mm"]))
    section.right_margin = Mm(int(FORMAT_RULES["margin_right_mm"]))
    section.header_distance = Mm(int(FORMAT_RULES["header_mm"]))
    section.footer_distance = Mm(int(FORMAT_RULES["footer_mm"]))
    _add_page_number_footer(section, note_font)

    # ── 표지 ────────────────────────────────────────────────
    _paragraph(document, f"{plan.year}년도 지방소멸대응기금 투자계획", body_font, 22, bold=True)
    _paragraph(document, f"충청북도 {plan.region}", body_font, 16)
    _paragraph(document, "", body_font, body_size)

    _paragraph(document, "이 문서는 초안입니다", note_font, note_size, bold=True, shade=MANUAL_SHADING)
    for line in (
        "데이터로 채울 수 있는 항목만 자동으로 채웠고, 나머지는 담당자가 작성해야 하는 구획으로 남겨 두었습니다.",
        f"회색 음영과 {MANUAL_MARKER} 표시가 있는 곳이 담당자 작성 구획입니다.",
        "실제 제출은 hwp 서식으로 변환해야 합니다. 제출 파일명과 서식은 안내서 p.3 을 따르십시오.",
        f"서식 출처: {TEMPLATE_SOURCE}",
        f"작성 기준: {plan.region} / {plan.year}년 / 계획서 ID {plan.plan_id} / 버전 {plan.version}",
    ):
        _paragraph(document, f"· {line}", note_font, note_size, shade=MANUAL_SHADING)
    for note in notes:
        _paragraph(document, f"· {note}", note_font, note_size, shade=MANUAL_SHADING)

    document.add_page_break()

    # ── 목차 ────────────────────────────────────────────────
    _paragraph(document, "목  차", body_font, 18, bold=True)
    for numeral, title in CHAPTERS.items():
        _paragraph(document, f"{numeral}. {title}", body_font, body_size)
    document.add_page_break()

    # ── 본문 ────────────────────────────────────────────────
    current_chapter = None
    for spec in SECTIONS:
        if spec.chapter != current_chapter:
            current_chapter = spec.chapter
            _paragraph(
                document,
                f"{spec.chapter}. {CHAPTERS[spec.chapter]}",
                body_font,
                18,
                bold=True,
                space_after=8,
            )

        _paragraph(document, f"{spec.number}. {spec.title}", body_font, 16, bold=True, space_after=6)

        stored = plan.sections.get(spec.section_id)
        content = (stored.content if stored else None) or ""

        if spec.fill_mode == "manual" or not content.strip():
            _paragraph(document, MANUAL_MARKER, note_font, note_size, bold=True, shade=MANUAL_SHADING)
            _paragraph(document, f"【작성내용】 {spec.writing_guide}", note_font, note_size, shade=MANUAL_SHADING)
            _paragraph(
                document,
                f"【기술 방향과 평가의 주안점】 {spec.evaluation_focus}",
                note_font,
                note_size,
                shade=MANUAL_SHADING,
            )
            if spec.reference_note:
                _paragraph(document, f"【참고사항】 {spec.reference_note}", note_font, note_size, shade=MANUAL_SHADING)
            _paragraph(document, f"(안내서 p.{spec.source_page})", note_font, note_size, shade=MANUAL_SHADING)
            _paragraph(document, "", body_font, body_size)

        if content.strip():
            for line in content.splitlines():
                _paragraph(document, line, body_font, body_size)
            if stored and stored.source == "human_input":
                _paragraph(document, "※ 위 내용은 담당자가 직접 입력한 값입니다.", note_font, note_size)
            if spec.manual_remainder:
                _paragraph(document, f"※ {spec.manual_remainder}", note_font, note_size, shade=MANUAL_SHADING)
            for warning in (stored.warnings if stored else []):
                _paragraph(document, f"※ 확인 필요: {warning}", note_font, note_size, shade=MANUAL_SHADING)

        _paragraph(document, "", body_font, body_size, space_after=10)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue(), notes


def build_pdf(plan: StoredPlan) -> tuple[bytes, list[str]]:
    """PDF 내보내기.

    한글 글꼴을 시스템에서 찾지 못해도 나오도록 reportlab 의 CID 글꼴을 쓴다.
    레이아웃은 docx 쪽이 원본이며, PDF 는 확인·공유용이다.
    """
    from reportlab.lib.enums import TA_LEFT
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer

    notes = [
        "PDF 는 확인·공유용이다. 제출용 서식(글꼴·여백)은 docx 를 hwp 로 변환해 맞춘다.",
    ]
    font_name = "HYSMyeongJo-Medium"
    try:
        pdfmetrics.registerFont(UnicodeCIDFont(font_name))
    except Exception as exc:  # pragma: no cover - 환경 의존
        logger.warning("PDF 한글 글꼴 등록 실패: %s", exc)
        notes.append("한글 글꼴 등록에 실패해 글자가 깨질 수 있다. docx 를 사용하는 것이 안전하다.")
        font_name = "Helvetica"

    buffer = io.BytesIO()
    document = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        topMargin=int(FORMAT_RULES["margin_top_mm"]) * mm,
        bottomMargin=int(FORMAT_RULES["margin_bottom_mm"]) * mm,
        leftMargin=int(FORMAT_RULES["margin_left_mm"]) * mm,
        rightMargin=int(FORMAT_RULES["margin_right_mm"]) * mm,
        title=f"{plan.year}년도 지방소멸대응기금 투자계획 초안 ({plan.region})",
    )
    body = ParagraphStyle("body", fontName=font_name, fontSize=11, leading=17, alignment=TA_LEFT)
    heading = ParagraphStyle("heading", parent=body, fontSize=15, leading=22, spaceBefore=10, spaceAfter=6)
    note = ParagraphStyle("note", parent=body, fontSize=9.5, leading=15, textColor="#444444")

    def escape(text: str) -> str:
        return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    flow = [
        Paragraph(escape(f"{plan.year}년도 지방소멸대응기금 투자계획"), heading),
        Paragraph(escape(f"충청북도 {plan.region}"), body),
        Spacer(1, 8),
        Paragraph(escape("이 문서는 초안입니다. 회색 구획은 담당자가 작성해야 합니다."), note),
        Paragraph(escape("실제 제출은 hwp 서식으로 변환해야 합니다."), note),
        Paragraph(escape(f"서식 출처: {TEMPLATE_SOURCE}"), note),
        PageBreak(),
    ]

    current_chapter = None
    for spec in SECTIONS:
        if spec.chapter != current_chapter:
            current_chapter = spec.chapter
            flow.append(Paragraph(escape(f"{spec.chapter}. {CHAPTERS[spec.chapter]}"), heading))
        flow.append(Paragraph(escape(f"{spec.number}. {spec.title}"), heading))

        stored = plan.sections.get(spec.section_id)
        content = (stored.content if stored else None) or ""
        if spec.fill_mode == "manual" or not content.strip():
            flow.append(Paragraph(escape(MANUAL_MARKER), note))
            flow.append(Paragraph(escape(f"【작성내용】 {spec.writing_guide}"), note))
            flow.append(Paragraph(escape(f"【기술 방향과 평가의 주안점】 {spec.evaluation_focus}"), note))
        if content.strip():
            for line in content.splitlines():
                if line.strip():
                    flow.append(Paragraph(escape(line), body))
        flow.append(Spacer(1, 6))

    document.build(flow)
    return buffer.getvalue(), notes
